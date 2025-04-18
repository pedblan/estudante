#!/usr/bin/env python3
# -*- coding: utf-8 -*-
#
# Copyright (c) 2024 Pedro Duarte Blanco
#
# Este software é distribuído sob a Licença MIT.
# Consulte o arquivo LICENSE para obter mais informações.

"""Constam deste script todas as funções que empregam IA para processar as demandas."""
import spacy
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from deep_translator import GoogleTranslator
from nltk.corpus import wordnet as wn

from openai import OpenAI
import whisper
import os
import tiktoken
from typing import Optional, Tuple

MODELO_ANALISE = "gpt-4o"
MODELO_TRANSCRICAO_API = "whisper-1"  # Empregado quando API==True
WHISPER_MODE = ['tiny', 'base', 'small', 'medium', 'large']  # Define o modelo Whisper como "base" por padrão

def get_openai_client() -> OpenAI:
    """Obtém o cliente da API OpenAI usando a chave da variável de ambiente.

    Returns:
        OpenAI: Cliente da API OpenAI.

    Raises:
        EnvironmentError: Se a chave da API não for encontrada nas variáveis de ambiente.
    """
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise EnvironmentError("A chave da API OpenAI não foi encontrada nas variáveis de ambiente.")
    return OpenAI(api_key=api_key)

def verificar_tokens(texto: str, modelo: str = MODELO_ANALISE, limite_tokens: int = 8192) -> bool:
    """Verifica se o número de tokens de um texto excede o limite permitido.

    Args:
        texto (str): Texto a ser verificado.
        modelo (str): Modelo de análise.
        limite_tokens (int): Limite máximo de tokens.

    Returns:
        bool: True se o número de tokens estiver dentro do limite, False caso contrário.
    """
    codificador = tiktoken.encoding_for_model(modelo)
    tokens = codificador.encode(texto)
    num_tokens = len(tokens)

    if num_tokens > limite_tokens:
        print(f"O número de tokens ({num_tokens}) excede o limite permitido ({limite_tokens}).")
        return False
    else:
        print(f"Sucesso: O número de tokens ({num_tokens}) está dentro do limite permitido.")
        return True

def resumir_texto(texto: str, modelo: str = MODELO_ANALISE, modo: str = 'geral', instrucao_personalizada: Optional[str] = None) -> str:
    """Recorre à API da OpenAI para resumir texto segundo instrução do usuário.

    Args:
        texto (str): Texto a ser resumido.
        modelo (str): Modelo de análise.
        modo (str): Modo de resumo ('geral', 'jurisprudencia', 'retorico' ou 'personalizado').
        instrucao_personalizada (Optional[str]): Instrução personalizada para o resumo.

    Returns:
        str: Texto resumido ou mensagem de erro.
    """
    client = get_openai_client()
    try:
        if modo == 'geral':
            instrucao = (
                "Você ajuda um analista de textos jornalísticos, políticos, jurídicos ou culturais. "
                "Sua função é a de sumarizar textos longos, destacando em tópicos os seguintes pontos:\n"
                "- principais pressupostos e premissas implícitas\n"
                "- pontos principais do argumento\n"
                "- detalhes e exemplos de cada ponto principal\n"
                "- identifique vulnerabilidades e possíveis vieses da argumentação."
            )
        elif modo == 'retorica':
            instrucao = ("Identifique as perguntas que o autor busca responder e como ele as responde."
                         "Caso o texto seja dividido em subtítulos, faça isso para cada subtítulo."
                         "Ao final, faça isso para o texto como um todo.")
        elif modo == 'jurisprudencia':
            instrucao = (
                "Explique a controvérsia jurídica em tela. Indique os argumentos empregados, "
                "ressaltando aquele que norteou o voto vencedor. Identifique vulnerabilidades nesse argumento."
            )
        elif modo == 'personalizado' and instrucao_personalizada:
            instrucao = instrucao_personalizada

        response = client.chat.completions.create(
            model=f"{modelo}",
            messages=[
                {"role": "system", "content": instrucao},
                {"role": "user", "content": f"Analise o seguinte texto: {texto}."}
            ],
            temperature=0.5,
            max_tokens=4000,
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0
        )
        return response.choices[0].message.content

    except Exception as e:
        return f"Ocorreu um erro ao tentar resumir o texto: {str(e)}"

def transcrever(arquivo_de_audio: str, idioma: str, api: bool) -> str:
    """Recebe um arquivo de áudio e realiza a transcrição usando API da OpenAI ou Whisper Local.

    Args:
        arquivo_de_audio (str): Caminho do arquivo de áudio.
        idioma (str): Idioma da transcrição.
        api (bool): Se True, usa a API da OpenAI; caso contrário, usa Whisper Local.

    Returns:
        str: Texto transcrito ou mensagem de erro.
    """
    client = get_openai_client()
    try:
        if api:
            with open(arquivo_de_audio, "rb") as audio_file:
                return client.audio.transcriptions.create(
                    model=MODELO_TRANSCRICAO_API,
                    file=audio_file,
                    language=idioma
                ).text
        else:
            model = whisper.load_model(WHISPER_MODE[1]) # seleciona o modelo 'base'
            return model.transcribe(arquivo_de_audio, language=idioma)['text']

    except Exception as e:
        return f"Erro ao transcrever o áudio: {str(e)}"

def ajustar_texto(texto: str, modelo: str = MODELO_ANALISE) -> str:
    """Ajusta o texto importado de imagens fotografadas para melhorar a legibilidade.

    Args:
        texto (str): Texto a ser ajustado.
        modelo (str): Modelo de análise.

    Returns:
        str: Texto ajustado ou mensagem de erro.
    """
    client = get_openai_client()
    try:
        instrucao = (
            "Você é um revisor de texto especializado em ajustar textos importados de imagens fotografadas.\n"
            "Os textos foram importados mediante OCR (tesseract).\n"
            "Você deve:\n"
            "- detectar o idioma em que o texto está escrito e ajustar erros decorrentes do mau reconhecimento de letras e palavras.\n"
            "- completar lacunas decorrentes do mau enquadramento do livro na fotografia, indicando com {completado} quando tiver feito isso\n"
            "- retornar o texto integral, com os ajustes."
        )

        response = client.chat.completions.create(
            model=f"{modelo}",
            messages=[
                {"role": "system", "content": instrucao},
                {"role": "user", "content": f"Ajuste a redação do seguinte texto: {texto}. Retorne todo o texto, com os devidos ajustes feitos."}
            ],
            temperature=1,
            max_tokens=8000,
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0
        )
        return response.choices[0].message.content

    except Exception as e:
        return f"Ocorreu um erro ao tentar ajustar o texto: {str(e)}"

def dividir_texto(texto: str) -> None:
    """Função para dividir texto (não implementada).

    Args:
        texto (str): Texto a ser dividido.
    """
    pass

def revisar(caminho_arquivo: str, idioma: str) -> Tuple[str, Document]:
    """Revisa um documento Word, destacando adjetivos, advérbios, verbos na voz passiva, conectores, verbos auxiliares e substantivos abstratos com cores distintas.

    Args:
        caminho_arquivo (str): Caminho do arquivo Word.
        idioma (str): Idioma do texto.

    Returns:
        Tuple[str, Document]: Título do documento e objeto Document revisado.
    """
    # Carregar o modelo spaCy grande para o idioma selecionado
    if idioma == 'pt':
        nlp = spacy.load('pt_core_news_lg')
    elif idioma == 'en':
        nlp = spacy.load('en_core_web_lg')
    else:
        raise ValueError("Idioma não suportado. Use 'pt' ou 'en'.")

    # Função para verificar se um verbo está na voz passiva
    def is_passive(verb):
        return 'Voice=Pass' in verb.morph

    # Função para aplicar destaque com cores diferentes
    def add_highlight(run, color):
        highlight = parse_xml(r'<w:highlight {} w:val="{}"/>'.format(nsdecls('w'), color))
        run._r.get_or_add_rPr().append(highlight)

    # Função para traduzir palavra para inglês, se necessário
    def traduzir_para_ingles(palavra: str) -> str:
        try:
            return GoogleTranslator(source='pt', target='en').translate(palavra)
        except Exception as e:
            print(f"Erro ao traduzir '{palavra}': {e}")
            return palavra

    # Função para verificar se uma palavra é substantivo abstrato
    def is_abstract(word: str, lang: str = 'por') -> bool:
        from nltk.corpus import wordnet as wn
        try:
            synsets = wn.synsets(word, lang=lang)
            if not synsets and lang == 'por':
                translated_word = traduzir_para_ingles(word)
                synsets = wn.synsets(translated_word, lang='eng')

            for syn in synsets:
                if syn.pos() == 'n' and "abstraction" in syn.lexname():
                    return True
            return False
        except Exception as e:
            print(f"Erro ao verificar substantivo abstrato para '{word}': {e}")
            return False

    # Mapear categorias para cores
    style_colors = {
        'ADJ': 'yellow',  # Adjetivos
        'ADV': 'green',   # Advérbios
        'VERB_PASSIVE': 'cyan',  # Verbos na voz passiva
        'CONNECTOR': 'pink',  # Conectores
        'AUX_VERB': 'lightblue',  # Verbos auxiliares
        'ABSTRACT_NOUN': 'orange'  # Substantivos abstratos
    }

    # Listas de conectores e verbos auxiliares para os idiomas
    conectores_pt = [
        "e", "ou", "mas", "porém", "portanto", "contudo", "todavia", "entretanto",
        "que", "no entanto", "além disso", "assim", "consequentemente", "nesse sentido"
    ]
    verbos_auxiliares_pt = ["ser", "estar", "ter", "haver"]
    conectores_en = [
        "and", "or", "but", "however", "therefore", "moreover", "furthermore",
        "that", "nevertheless", "besides", "thus", "consequently"
    ]
    verbos_auxiliares_en = ["be", "have", "do", "shall", "will", "may", "can"]

    conectores = conectores_pt if idioma == 'pt' else conectores_en
    verbos_auxiliares = verbos_auxiliares_pt if idioma == 'pt' else verbos_auxiliares_en

    # Carregar o documento Word
    doc = Document(caminho_arquivo)

    for para in doc.paragraphs:
        doc_para = nlp(para.text)

        # Reescrever o parágrafo diretamente preservando a formatação
        new_paragraph = []

        for token in doc_para:
            if token.pos_ == 'ADJ':
                new_paragraph.append((token.text, style_colors['ADJ']))
            elif token.pos_ == 'ADV':
                new_paragraph.append((token.text, style_colors['ADV']))
            elif token.pos_ == 'VERB' and is_passive(token):
                new_paragraph.append((token.text, style_colors['VERB_PASSIVE']))
            elif token.lemma_.lower() in conectores:  # Usa o lema para conectores
                new_paragraph.append((token.text, style_colors['CONNECTOR']))
            elif token.lemma_.lower() in verbos_auxiliares:  # Usa o lema para verbos auxiliares
                new_paragraph.append((token.text, style_colors['AUX_VERB']))
            elif token.pos_ == 'NOUN' and is_abstract(token.text, lang='por'):
                new_paragraph.append((token.text, style_colors['ABSTRACT_NOUN']))
            else:
                new_paragraph.append((token.text, None))

            if token.whitespace_:
                new_paragraph.append((token.whitespace_, None))

        # Limpar e reconstruir o parágrafo
        para.clear()
        for text, color in new_paragraph:
            run = para.add_run(text)
            if color:
                add_highlight(run, color)

    # Extrair o título do arquivo
    titulo = os.path.splitext(os.path.basename(caminho_arquivo))[0]

    return titulo, doc
