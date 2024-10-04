#!/usr/bin/env python3
# -*- coding: utf-8 -*-
#
# Copyright (c) 2024 Pedro Duarte Blanco
#
# Este software é distribuído sob a Licença MIT.
# Consulte o arquivo LICENSE para obter mais informações.


import yt_dlp
import os
import math
import ffmpeg
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
import re
import fitz
from proc import *
from pathlib import Path
import re
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import subprocess



# Definir a pasta temporária e o caminho de destino (Desktop)
TEMP_FOLDER = 'temp'
os.makedirs(TEMP_FOLDER, exist_ok=True)  # Cria a pasta 'temp' se não existir
PASTA_DESTINO = os.path.join(os.path.expanduser('~'), 'Desktop')  # Desktop como padrão

def download_yt(youtube_url):
    """Extrai áudio de vídeo de streaming e o grava em temp. Retorna o título do vídeo e o arquivo de áudio temporário."""
    try:
        ydl_opts = {
            'format': 'bestaudio/best',
            'postprocessors': [{
                'key': 'FFmpegExtractAudio',
                'preferredcodec': 'mp3',
                'preferredquality': '192'
            }],
            'outtmpl': os.path.join(TEMP_FOLDER, 'temp.%(ext)s')  # Salva na pasta temp com o título do vídeo
        }

        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            print("Baixando vídeo...")
            info_dict = ydl.extract_info(youtube_url, download=True)
            title = info_dict.get('title')
            ext = info_dict.get('ext', 'mp3')  # Assume extensão mp3 como padrão

        print("Extraindo áudio do vídeo...")
        file_path = f"{TEMP_FOLDER}/temp.mp3"

        return title, file_path

    except Exception as e:
        return None, f"Ocorreu um erro ao baixar o áudio do streaming: {str(e)}"


def dividir_audio(caminho_audio='temp/temp.mp3', chunk_size_mb=8):
    """Divide um arquivo de áudio. Retorna o título e uma lista com o caminho de cada parte."""
    print("Analisando áudio...")
    try:
        if not os.path.exists(caminho_audio):
            raise FileNotFoundError(f"Arquivo de áudio não encontrado: {caminho_audio}")

        titulo = os.path.splitext(os.path.basename(caminho_audio))[0]
        chunk_size_bytes = chunk_size_mb * 1024 * 1024

        # Tenta obter informações sobre o arquivo de áudio
        probe = ffmpeg.probe(caminho_audio)
        duration = float(probe['format']['duration'])
        file_size = int(probe['format']['size'])

        # Calcula o número de chunks com base na duração
        num_chunks = math.ceil(file_size / chunk_size_bytes)  # Garante que o último pedaço será incluído
        chunk_duration = duration / num_chunks  # Divide a duração igualmente pelos chunks

        chunk_paths = []
        for i in range(num_chunks):
            chunk_file = os.path.join(TEMP_FOLDER, f"chunk_{i}{os.path.splitext(caminho_audio)[1]}")

            try:
                ffmpeg.input(caminho_audio, ss=i * chunk_duration, t=chunk_duration).output(chunk_file).run(
                    overwrite_output=True)
                chunk_paths.append(chunk_file)
            except ffmpeg.Error as e:
                print(f"Erro ao processar o chunk {i}: {str(e)}")

        return titulo, chunk_paths, duration

    except Exception as e:
        print(f"Erro ao dividir o áudio: {str(e)}")
        return None, []

def gravar_documento(titulo, doc):
    print("Salvando documento...")
    """Grava o documento Word no desktop. Retorna o caminho do arquivo ou uma mensagem de erro."""
    try:
        file_path = normalizar_nome_do_arquivo(titulo)
        doc.save(file_path)
        print(f"Documento salvo no desktop: {file_path}")
        return file_path
    except Exception as e:
        print(f"Erro ao tentar salvar o documento: {str(e)}")
        return f"Erro ao tentar salvar o documento: {str(e)}"

def limpar_temp():
    """Limpa a pasta temporária."""
    print("Apagando arquivos temporários...")
    try:
        for filename in os.listdir(TEMP_FOLDER):
            file_path = os.path.join(TEMP_FOLDER, filename)
            if os.path.isfile(file_path):
                os.remove(file_path)
        print("Pasta temporária limpa com sucesso.")
    except Exception as e:
        print(f"Erro ao limpar a pasta 'temp': {str(e)}")

def extrair_texto(caminho_arquivo):
    print("Extraindo texto do arquivo...")
    """Extrai texto de arquivos PDF e Word. Retorna string."""
    try:
        # Verifica a extensão do arquivo
        extensao = os.path.splitext(caminho_arquivo)[1].lower()
        titulo = os.path.splitext(caminho_arquivo)[0]

        if extensao == '.pdf':
            # Extração de texto para arquivos PDF
            texto = ""
            with fitz.open(caminho_arquivo) as doc:
                for pagina in doc:
                    texto += pagina.get_text()
            return titulo, texto

        elif extensao == '.docx':
            # Extração de texto para arquivos Word (.docx)
            doc = Document(caminho_arquivo)
            texto = "\n".join([paragrafo.text for paragrafo in doc.paragraphs])
            return titulo, texto

        else:
            return f"Formato de arquivo {extensao} não suportado. Use PDF ou DOCX."

    except Exception as e:
        return f"Erro ao extrair texto: {str(e)}"


def abrir_doc_produzido(caminho_arquivo):
    print("Abrindo documento...")
    """Abre o arquivo Word salvo no desktop. Utiliza tratamento de erro. Retorna True ou None (erro)."""
    try:
        if os.path.exists(caminho_arquivo):
            if os.name == 'nt':  # Para Windows
                # Usa o subprocess para chamar o start do Windows de forma mais segura
                subprocess.run(['start', caminho_arquivo], shell=True)
            elif os.name == 'posix':  # Para Mac/Linux
                # Usa subprocess para abrir arquivos de forma segura, substituindo os.system
                subprocess.run(['open', caminho_arquivo], check=True)
            print(f"Abrindo o arquivo Word: {caminho_arquivo}")
            return True
        else:
            print(f"Arquivo não encontrado: {caminho_arquivo}")
            return None

    except Exception as e:
        print(f"Erro ao tentar abrir o arquivo: {str(e)}")
        return None



def gravar_audio():
    """Função para gravar áudio (não implementada)"""
    pass

def normalizar_nome_do_arquivo(titulo):
    """Retorna caminho do arquivo normalizado, removendo caracteres problemáticos e garantindo que o nome seja seguro."""
    # Substitui caracteres inválidos por underscores e garante que o título seja seguro
    titulo = re.sub(r'[ \\/*?:"<>|()]+', "_", titulo)
    return os.path.join(PASTA_DESTINO, f"{titulo}.docx")

def adicionar_ao_word(doc, trecho_transcrito, max_palavras=80):
    """Adiciona o trecho transcrito ao documento Word, quebrando a cada aproximadamente `max_palavras` palavras,
    sem dividir no meio de uma frase. Retorna o objeto Document."""

    try:
        # Dividimos o texto em frases com base nos terminais de sentença (. ! ?)
        frases = re.split(r'(?<=[.!?]) +', trecho_transcrito.strip())

        bloco_palavras = []  # Para armazenar temporariamente o bloco de frases
        contador_palavras = 0

        for frase in frases:
            palavras_frase = frase.split()
            bloco_palavras.append(frase)
            contador_palavras += len(palavras_frase)

            # Quando o número de palavras do bloco chega a `max_palavras` ou mais, adiciona ao Word
            if contador_palavras >= max_palavras:
                paragrafo_texto = ' '.join(bloco_palavras)
                paragrafo = doc.add_paragraph(paragrafo_texto)
                paragrafo.paragraph_format.first_line_indent = Cm(1.5)  # Indentar o texto principal
                paragrafo.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justificar o texto

                # Reseta o bloco de palavras e contador
                bloco_palavras = []
                contador_palavras = 0

        # Adiciona o bloco restante, caso tenha sobrado alguma frase
        if bloco_palavras:
            paragrafo_texto = ' '.join(bloco_palavras)
            paragrafo = doc.add_paragraph(paragrafo_texto)
            paragrafo.paragraph_format.first_line_indent = Cm(1.5)
            paragrafo.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        return doc

    except Exception as e:
        print(f"Erro ao adicionar texto ao documento: {str(e)}")
        return None

def paragrafo_timestamp(tempo_atual, doc, trecho_transcrito, max_palavras=80):
    """Cria um parágrafo no Word com o timestamp na margem esquerda e adiciona o texto fornecido, quebrado em blocos."""
    try:
        # Calcular o timestamp (em minutos e segundos) e arredondar os segundos
        minutos = int(tempo_atual // 60)
        segundos = int(round(tempo_atual % 60))  # Arredondar os segundos

        timestamp = f"[{minutos:02}:{segundos:02}]"

        # Criar o parágrafo de timestamp com alinhamento à esquerda
        paragrafo_timestamp = doc.add_paragraph(timestamp)
        paragrafo_timestamp.paragraph_format.left_indent = -Cm(1.5)  # Ajuste da indentação negativa
        paragrafo_timestamp.runs[0].font.size = Pt(10)  # Tamanho de fonte para o timestamp

        # Adicionar o texto transcrito com quebra a cada 80 palavras
        adicionar_ao_word(doc, trecho_transcrito, max_palavras)

        return doc
    except Exception as e:
        print(f"Erro ao adicionar timestamp e texto ao documento: {str(e)}")
        return None

def transcrever_partes(lista_de_partes, idioma, api, duracao_total, max_palavras, com_timestamp):
    """Transcreve partes do áudio ou vídeo com ou sem timestamps, dividindo o texto em blocos. Retorna objeto doc."""
    doc = Document()

    # Duração total em segundos
    duracao_total_segundos = duracao_total

    # Calcular a duração de cada parte dividida
    duracao_por_parte = duracao_total_segundos / len(lista_de_partes)

    # Tempo atual no vídeo
    tempo_atual = 0

    for i, parte in enumerate(lista_de_partes):
        try:
            # Transcrever a parte
            trecho_transcrito = transcrever(parte, idioma, api)

            if trecho_transcrito:
                if com_timestamp:
                    # Adicionar o timestamp e o texto ao documento
                    paragrafo_timestamp(tempo_atual, doc, trecho_transcrito, max_palavras)
                else:
                    # Adicionar o texto ao documento sem timestamp
                    adicionar_ao_word(doc, trecho_transcrito, max_palavras)

                # Atualizar o tempo atual para o próximo trecho
                tempo_atual += duracao_por_parte
            else:
                print(f"Parte {i} não foi transcrita corretamente.")
        except Exception as e:
            print(f"Erro ao transcrever a parte {i}: {str(e)}")

    return doc


def adicionar_com_subtitulos(doc, resumo):
    """
    Adiciona o conteúdo do resumo ao documento Word com formatação específica para títulos, subtítulos e parágrafos.
    Títulos em negrito e subtítulos sublinhados com os números na mesma linha.
    """
    for linha in resumo.splitlines():
        if linha.startswith('###'):
            # Formatando títulos (negrito)
            linha_formatada = linha.replace('### ', '')  # Remove o marcador ### do início
            paragrafo = doc.add_paragraph(linha_formatada)
            run = paragrafo.runs[0]
            run.bold = True
            run.font.size = Pt(14)  # Tamanho maior para títulos
        elif linha.startswith('**'):
            # Formatando subtítulos (sublinhado) e número na mesma linha
            linha_formatada = linha.replace('**', '').strip()  # Remove os asteriscos e espaços extras
            paragrafo = doc.add_paragraph()
            run = paragrafo.add_run(linha_formatada)
            run.underline = True
            run.font.size = Pt(12)  # Tamanho menor para subtítulos
        elif linha.startswith('-'):
            # Formatando as vulnerabilidades, sem sublinhado ou negrito
            paragrafo = doc.add_paragraph(linha.replace('-', '').strip())
        elif linha:
            # Mantendo o parágrafo normal, sem formatação específica
            doc.add_paragraph(linha.strip())
    return doc