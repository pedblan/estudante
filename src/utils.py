#!/usr/bin/env python3
# -*- coding: utf-8 -*-
#
# Copyright (c) 2024 Pedro Duarte Blanco
#
# Este software é distribuído sob a Licença MIT.
# Consulte o arquivo LICENSE para obter mais informações.

"""Este script contém funções que são empregadas nas tarefas principais do aplicativo."""

import io
import math
import os
import re
import subprocess
import sys
import warnings
from pathlib import Path
from typing import List, Tuple, Union

import ffmpeg
import fitz
import pytesseract
import yt_dlp
from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from src.paths import OUTPUT_DIR, TEMP_DIR
from src.utils_ia import ajustar_texto, transcrever

# Pastas temporária e de saída
TEMP_FOLDER = TEMP_DIR
OUTPUT_FOLDER = OUTPUT_DIR


def download_yt(youtube_url: str, temp_folder: Path = TEMP_FOLDER) -> Tuple[Union[str, None], str]:
    """Extrai áudio de vídeo de streaming e o grava em temp.

    Args:
        youtube_url (str): URL do vídeo do YouTube.

    Returns:
        tuple: Título do vídeo e o arquivo de áudio temporário.
    """
    try:
        ydl_opts = {
            "format": "bestaudio/best",
            "postprocessors": [
                {
                    "key": "FFmpegExtractAudio",
                    "preferredcodec": "mp3",
                    "preferredquality": "192",
                }
            ],
            "outtmpl": str(temp_folder / "temp.%(ext)s"),
        }

        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(youtube_url, download=True)
            title = info.get("title")

        file_path = str(temp_folder / "temp.mp3")
        return title, file_path

    except Exception as e:
        return None, str(e)


def dividir_audio(
    caminho_audio: str,
    chunk_size_mb: int = 8,
    temp_folder: Path = TEMP_FOLDER,
) -> Tuple[Union[str, None], List[str], Union[float, None]]:
    """Divide um arquivo de áudio.

    Args:
        caminho_audio (str): Caminho do arquivo de áudio.
        chunk_size_mb (int): Tamanho de cada parte em MB.

    Returns:
        tuple: Título e uma lista com o caminho de cada parte.
    """
    try:
        if not os.path.exists(caminho_audio):
            raise FileNotFoundError(f"Arquivo de áudio não encontrado: {caminho_audio}")

        titulo = Path(caminho_audio).stem
        chunk_size_bytes = chunk_size_mb * 1024 * 1024

        # Tenta obter informações sobre o arquivo de áudio
        probe = ffmpeg.probe(caminho_audio)
        duration = float(probe['format']['duration'])
        file_size = int(probe['format']['size'])

        # Calcula o número de chunks com base na duração
        num_chunks = math.ceil(file_size / chunk_size_bytes)  # Garante que o último pedaço será incluído
        chunk_duration = duration / num_chunks  # Divide a duração igualmente pelos chunks

        chunk_paths: List[str] = []
        for i in range(num_chunks):
            chunk_file = temp_folder / f"chunk_{i}{os.path.splitext(caminho_audio)[1]}"

            try:
                ffmpeg.input(caminho_audio, ss=i * chunk_duration, t=chunk_duration).output(str(chunk_file)).run(
                    overwrite_output=True)
                chunk_paths.append(str(chunk_file))
            except ffmpeg.Error:
                continue

        return titulo, chunk_paths, duration

    except Exception:
        return None, [], None


def gravar_documento(titulo: str, doc: Document, output_folder: Path = OUTPUT_FOLDER) -> str:
    """Grava o documento Word no desktop.

    Args:
        titulo (str): Título do documento.
        doc (Document): Objeto Document do python-docx.

    Returns:
        str: Caminho do arquivo ou uma mensagem de erro.
    """
    try:
        file_path = normalizar_nome_do_arquivo(titulo, output_folder)
        doc.save(file_path)
        return file_path
    except Exception as e:
        return f"Erro ao tentar salvar o documento: {str(e)}"


def limpar_temp(temp_folder: Path = TEMP_FOLDER) -> None:
    """Limpa a pasta temporária."""
    try:
        for filename in os.listdir(temp_folder):
            file_path = temp_folder / filename
            if file_path.is_file():
                file_path.unlink()
    except Exception:
        pass


def extrair_texto(caminho_arquivo: str) -> Tuple[str, str]:
    """Extrai texto de arquivos PDF e Word.

    Args:
        caminho_arquivo (str): Caminho do arquivo.

    Returns:
        tuple: Título e texto extraído.
    """
    try:
        # Verifica a extensão do arquivo
        caminho = Path(caminho_arquivo)
        extensao = caminho.suffix.lower()
        titulo = str(caminho.with_suffix(''))

        if extensao == '.pdf':
            # Extração de texto para arquivos PDF
            texto = ""
            with fitz.open(caminho_arquivo) as doc:
                for pagina in doc:
                    texto += pagina.get_text()
            return titulo, texto

        elif extensao == '.docx':
            # Extração de texto para arquivos Word (.docx)
            doc = Document(caminho_arquivo)
            texto = "\n".join([paragrafo.text for paragrafo in doc.paragraphs])
            return titulo, texto

        else:
            return titulo, f"Formato de arquivo {extensao} não suportado. Use PDF ou DOCX."

    except Exception as e:
        return "", f"Erro ao extrair texto: {str(e)}"


def abrir_doc_produzido(caminho_arquivo: str) -> Union[bool, None]:
    """Abre o arquivo Word salvo no desktop.

    Args:
        caminho_arquivo (str): Caminho do arquivo.

    Returns:
        bool: True se o arquivo foi aberto com sucesso, None caso contrário.
    """
    try:
        caminho = Path(caminho_arquivo)
        if caminho.exists():
            if os.name == 'nt':  # Para Windows
                # Usa o subprocess para chamar o start do Windows de forma mais segura
                subprocess.run(['start', str(caminho)], shell=True)
            elif os.name == 'posix':  # Para Mac/Linux
                # Usa subprocess para abrir arquivos de forma segura, substituindo os.system
                subprocess.run(['open', str(caminho)], check=True)
            return True
        else:
            return None

    except Exception:
        return None


def gravar_audio() -> None:
    """Função para gravar áudio (não implementada)."""
    pass


def normalizar_nome_do_arquivo(titulo: str, output_folder: Path = OUTPUT_FOLDER) -> str:
    """Retorna caminho do arquivo normalizado, removendo caracteres problemáticos e garantindo que o nome seja seguro.

    Args:
        titulo (str): Título do arquivo.

    Returns:
        str: Caminho do arquivo normalizado.
    """
    # Substitui caracteres inválidos por underscores e garante que o título seja seguro
    titulo = re.sub(r'[ \\/*?:"<>|()]+', "_", titulo)
    return str(output_folder / f"{titulo}.docx")


def adicionar_ao_word(doc: Document, trecho_transcrito: str, max_palavras: int = 80) -> Union[Document, None]:
    """Adiciona o trecho transcrito ao documento Word, quebrando a cada aproximadamente `max_palavras` palavras,
    sem dividir no meio de uma frase.

    Args:
        doc (Document): Objeto Document do python-docx.
        trecho_transcrito (str): Texto transcrito.
        max_palavras (int): Número máximo de palavras por parágrafo.

    Returns:
        Document: Objeto Document atualizado.
    """
    try:
        # Dividimos o texto em frases com base nos terminais de sentença (. ! ?)
        frases = re.split(r'(?<=[.!?]) +', trecho_transcrito.strip())

        bloco_palavras = []  # Para armazenar temporariamente o bloco de frases
        contador_palavras = 0

        for frase in frases:
            palavras_frase = frase.split()
            bloco_palavras.append(frase)
            contador_palavras += len(palavras_frase)

            # Quando o número de palavras do bloco chega a `max_palavras` ou mais, adiciona ao Word
            if contador_palavras >= max_palavras:
                paragrafo_texto = ' '.join(bloco_palavras)
                paragrafo = doc.add_paragraph(paragrafo_texto)
                paragrafo.paragraph_format.first_line_indent = Cm(1.5)  # Indentar o texto principal
                paragrafo.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justificar o texto

                # Reseta o bloco de palavras e contador
                bloco_palavras = []
                contador_palavras = 0

        # Adiciona o bloco restante, caso tenha sobrado alguma frase
        if bloco_palavras:
            paragrafo_texto = ' '.join(bloco_palavras)
            paragrafo = doc.add_paragraph(paragrafo_texto)
            paragrafo.paragraph_format.first_line_indent = Cm(1.5)
            paragrafo.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        return doc

    except Exception as e:
        print(f"Erro ao adicionar texto ao documento: {str(e)}")
        return None


def paragrafo_timestamp(tempo_atual: float, doc: Document, trecho_transcrito: str, max_palavras: int = 80) -> Union[Document, None]:
    """Cria um parágrafo no Word com o timestamp na margem esquerda e adiciona o texto fornecido, quebrado em blocos.

    Args:
        tempo_atual (float): Tempo atual em segundos.
        doc (Document): Objeto Document do python-docx.
        trecho_transcrito (str): Texto transcrito.
        max_palavras (int): Número máximo de palavras por parágrafo.

    Returns:
        Document: Objeto Document atualizado.
    """
    try:
        # Calcular o timestamp (em minutos e segundos) e arredondar os segundos
        minutos = int(tempo_atual // 60)
        segundos = int(round(tempo_atual % 60))  # Arredondar os segundos

        timestamp = f"[{minutos:02}:{segundos:02}]"

        # Criar o parágrafo de timestamp com alinhamento à esquerda
        paragrafo_timestamp = doc.add_paragraph(timestamp)
        paragrafo_timestamp.paragraph_format.left_indent = -Cm(1.5)  # Ajuste da indentação negativa
        paragrafo_timestamp.runs[0].font.size = Pt(10)  # Tamanho de fonte para o timestamp

        # Adicionar o texto transcrito com quebra a cada 80 palavras
        adicionar_ao_word(doc, trecho_transcrito, max_palavras)

        return doc
    except Exception as e:
        print(f"Erro ao adicionar timestamp e texto ao documento: {str(e)}")
        return None


def transcrever_partes(lista_de_partes: List[str], idioma: str, api: bool, duracao_total: float, max_palavras: int, com_timestamp: bool) -> Document:
    """Transcreve partes do áudio ou vídeo com ou sem timestamps, dividindo o texto em blocos.

    Args:
        lista_de_partes (list): Lista de caminhos das partes do áudio/vídeo.
        idioma (str): Idioma da transcrição.
        api (bool): Se True, usa a API da OpenAI; caso contrário, usa Whisper Local.
        duracao_total (float): Duração total do áudio/vídeo em segundos.
        max_palavras (int): Número máximo de palavras por parágrafo.
        com_timestamp (bool): Se True, adiciona timestamps.

    Returns:
        Document: Objeto Document com a transcrição.
    """
    doc = Document()

    # Duração total em segundos
    duracao_total_segundos = duracao_total

    # Calcular a duração de cada parte dividida
    duracao_por_parte = duracao_total_segundos / len(lista_de_partes)

    # Tempo atual no vídeo
    tempo_atual = 0

    for i, parte in enumerate(lista_de_partes):
        try:
            # Transcrever a parte
            trecho_transcrito = transcrever(parte, idioma, api)

            if trecho_transcrito:
                if com_timestamp:
                    # Adicionar o timestamp e o texto ao documento
                    paragrafo_timestamp(tempo_atual, doc, trecho_transcrito, max_palavras)
                else:
                    # Adicionar o texto ao documento sem timestamp
                    adicionar_ao_word(doc, trecho_transcrito, max_palavras)

                # Atualizar o tempo atual para o próximo trecho
                tempo_atual += duracao_por_parte
            else:
                pass
        except Exception:
            pass

    return doc


def adicionar_com_subtitulos(doc: Document, resumo: str) -> Document:
    """Adiciona o conteúdo do resumo ao documento Word com formatação específica para títulos, subtítulos e parágrafos.

    Args:
        doc (Document): Objeto Document do python-docx.
        resumo (str): Texto do resumo.

    Returns:
        Document: Objeto Document atualizado.
    """
    for linha in resumo.splitlines():
        if linha.startswith('###'):
            # Formatando títulos (negrito)
            linha_formatada = linha.replace('### ', '')  # Remove o marcador ### do início
            paragrafo = doc.add_paragraph(linha_formatada)
            run = paragrafo.runs[0]
            run.bold = True
            run.font.size = Pt(14)  # Tamanho maior para títulos
        elif linha.startswith('**'):
            # Formatando subtítulos (sublinhado) e número na mesma linha
            linha_formatada = linha.replace('**', '').strip()  # Remove os asteriscos e espaços extras
            paragrafo = doc.add_paragraph()
            run = paragrafo.add_run(linha_formatada)
            run.underline = True
            run.font.size = Pt(12)  # Tamanho menor para subtítulos
        elif linha.startswith('-'):
            # Formatando as vulnerabilidades, sem sublinhado ou negrito
            paragrafo = doc.add_paragraph(linha.replace('-', '').strip())
        elif linha:
            # Mantendo o parágrafo normal, sem formatação específica
            doc.add_paragraph(linha.strip())
    return doc


def reconhecer_ocr(caminho_arquivo: str, ajustar_com_api: bool) -> Tuple[str, str]:
    """Reconhece texto em imagens de um arquivo PDF usando OCR.

    Args:
        caminho_arquivo (str): Caminho do arquivo PDF.

    Returns:
        tuple: Título e texto reconhecido.
    """
    # Abrir o PDF
    documento = fitz.open(caminho_arquivo)
    texto_revisado = ""

    for num_pagina in range(len(documento)):
        pagina = documento.load_page(num_pagina)
        pix = pagina.get_pixmap()  # Extrai a imagem da página
        imagem = Image.open(io.BytesIO(pix.tobytes("png")))  # Converte em imagem PIL

        # Usa o Tesseract para extrair o texto da imagem
        texto_ocr = pytesseract.image_to_string(imagem)
        if ajustar_com_api:
            texto_ocr_ajustado = ajustar_texto(texto_ocr)
            texto_revisado += texto_ocr_ajustado
        else:
            texto_revisado += texto_ocr

    documento.close()
    titulo = Path(caminho_arquivo).stem

    return titulo, texto_revisado

def suprimir_avisos(opcao_log_simplificado=True) -> None:
    """Suprime avisos e redireciona o fluxo de erro padrão para null.

    Esta função suprime avisos das categorias FutureWarning e UserWarning,
    e redireciona o fluxo de erro padrão (stderr) para null, de modo que
    apenas as declarações print sejam exibidas no console.

    Returns:
        None: Esta função não retorna nenhum valor.
    """
    if opcao_log_simplificado:
        # Suprimir avisos específicos
        warnings.filterwarnings("ignore", category=FutureWarning)
        warnings.filterwarnings("ignore", category=UserWarning)

        # Redirecionar o fluxo de erro padrão para null
        sys.stderr = open(os.devnull, 'w')

        # Suprimir avisos do ffmpeg
        ffmpeg_log_level = 'quiet'
        os.environ['FFMPEG_LOG_LEVEL'] = ffmpeg_log_level
